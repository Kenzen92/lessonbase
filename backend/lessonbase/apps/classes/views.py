import datetime

from django.db.models import Count, Sum, Q, Case, When, Value, Avg
from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from openai import OpenAI
from rest_framework.decorators import (
    api_view,
    permission_classes,
    authentication_classes,
    action,
)
from apps.core.authentication import ExpiringTokenAuthentication as TokenAuthentication
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.viewsets import ModelViewSet
from rest_framework import status
from apps.user_accounts.models import ClassGroup, Student, Teacher
from apps.classes.serialisers import (
    ClassEventSerializer,
    ClassEventCreateSerializer,
    ClassEventDateOrderedSerializer,
    SessionFeedbackSerializer,
)
from apps.assignments.models import Assignment
from .models import ClassEvent, SessionFeedback
from rest_framework import viewsets
from django.utils import timezone
from datetime import timedelta
from django.db import IntegrityError, transaction


class ClassEventViewSet(viewsets.ViewSet):
    authentication_classes = [TokenAuthentication]
    permission_classes = [IsAuthenticated]

    def get_serializer_class(self):
        if self.action == "create":
            return ClassEventCreateSerializer
        if self.action == "list":
            return ClassEventDateOrderedSerializer
        return ClassEventSerializer

    def list(self, request):
        user = request.user.get_real_instance()
        if isinstance(user, Teacher):
            class_events = (
                ClassEvent.objects.filter(teachers=user, classroom_type="scheduled")
                .distinct()
                .select_related("subject")
                .prefetch_related("students")
            )
        else:
            class_events = (
                ClassEvent.objects.filter(students=user, classroom_type="scheduled")
                .distinct()
                .select_related("subject")
                .prefetch_related("teachers")
            )
        serializer_class = self.get_serializer_class()
        serializer = serializer_class(class_events, many=True)
        return Response(serializer.data)

    def create(self, request):
        teacher_ids = [request.user.pk]
        student_ids = request.data.get("students", [])
        teachers = Teacher.objects.filter(pk__in=teacher_ids)
        students = Student.objects.filter(pk__in=student_ids)
        serializer_class = self.get_serializer_class()
        serializer = serializer_class(data=request.data)
        if serializer.is_valid():
            class_event = serializer.save()
            class_event.teachers.set(teachers)
            class_event.students.set(students)
            return Response(
                {"message": "Class event created successfully"},
                status=status.HTTP_201_CREATED,
            )
        return Response({"error": serializer.errors}, status=status.HTTP_400_BAD_REQUEST)

    def retrieve(self, request, pk=None):
        try:
            class_event = ClassEvent.objects.get(id=pk)
            serializer = ClassEventSerializer(class_event)
            return Response(serializer.data)
        except ClassEvent.DoesNotExist:
            return Response(
                {"error": "Class event not found"}, status=status.HTTP_404_NOT_FOUND
            )

    def update(self, request, pk=None):
        return self._update_class_event(request, pk, partial=False)

    def partial_update(self, request, pk=None):
        return self._update_class_event(request, pk, partial=True)

    def destroy(self, request, pk=None):
        try:
            class_event = ClassEvent.objects.get(id=pk)
            class_event.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except ClassEvent.DoesNotExist:
            return Response(
                {"error": "Class event not found"}, status=status.HTTP_404_NOT_FOUND
            )

    def _update_class_event(self, request, pk, partial):
        if not pk:
            return Response(
                {"error": "Class event ID is required"},
                status=status.HTTP_400_BAD_REQUEST,
            )
        try:
            class_event = ClassEvent.objects.get(id=pk)
        except ClassEvent.DoesNotExist:
            return Response(
                {"error": "Class event not found"}, status=status.HTTP_404_NOT_FOUND
            )
        teacher_ids = [request.user.pk]
        student_ids = request.data.get("students", [])
        teachers = Teacher.objects.filter(pk__in=teacher_ids)
        students = Student.objects.filter(pk__in=student_ids)
        serializer_class = self.get_serializer_class()
        serializer = serializer_class(class_event, data=request.data, partial=partial)
        if serializer.is_valid():
            class_event = serializer.save()
            class_event.teachers.set(teachers)
            class_event.students.set(students)
            return Response({"message": "Class event updated successfully"})
        return Response({"error": serializer.errors}, status=status.HTTP_400_BAD_REQUEST)


@api_view(["GET", "POST", "DELETE", "PUT", "PATCH"])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def class_events_for_student(request, student_id=None):
    if request.method == "GET":
        class_events = ClassEvent.objects.filter(
            students=student_id, classroom_type="scheduled"
        )
        serializer = ClassEventSerializer(class_events, many=True)
        return Response(serializer.data)


@api_view(["GET"])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def student_statistics(request):
    from apps.resources.models import Resource

    page = request.GET.get("page")
    try:
        student = Student.objects.get(pk=request.user.id)
        current_datetime = datetime.datetime.now()

        total_classes = ClassEvent.objects.filter(
            students=student, classroom_type="scheduled"
        ).count()
        completed_classes = ClassEvent.objects.filter(
            students=student,
            classroom_type="scheduled",
            start_time__lt=current_datetime,
        ).count()
        upcoming_classes = ClassEvent.objects.filter(
            students=student,
            classroom_type="scheduled",
            start_time__gte=current_datetime,
        ).count()
        total_assignments = Assignment.objects.filter(students=student).count()

        if page == "dashboard":
            stats = {
                "total_classes": total_classes,
                "completed_classes": completed_classes,
                "upcoming_classes": upcoming_classes,
            }
        elif page == "assignments":
            stats = {
                "total_assignments": total_assignments,
                "total_documents": Resource.objects.filter(
                    assignment_links__assignment__students=student
                )
                .distinct()
                .count(),
            }
        else:
            return Response(
                {"error": "Invalid page parameter"}, status=status.HTTP_400_BAD_REQUEST
            )

        return Response({"data": stats}, status=status.HTTP_200_OK)

    except Student.DoesNotExist:
        return Response(
            {"error": "Student not found"}, status=status.HTTP_404_NOT_FOUND
        )


@api_view(["GET"])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def teacher_statistics(request):
    from apps.resources.models import Resource

    try:
        teacher = Teacher.objects.get(pk=request.user.id)
        current_datetime = datetime.datetime.now()

        total_students = teacher.students.count()
        upcoming_classes = ClassEvent.objects.filter(
            teachers=teacher,
            classroom_type="scheduled",
            start_time__gte=current_datetime,
        ).count()
        total_assignments = Assignment.objects.filter(teachers=teacher).count()
        total_class_groups = ClassGroup.objects.filter(teachers=teacher).count()
        average_students_per_group = (
            ClassGroup.objects.filter(teachers=teacher)
            .annotate(num_students=Count("students"))
            .aggregate(avg_students=Avg("num_students"))
        )

        stats = {
            "total_students": total_students,
            "upcoming_classes": upcoming_classes,
            "pending_assignments": Assignment.objects.filter(
                teachers=teacher, marked=False
            ).count(),
            "total_teaching_hours": ClassEvent.objects.filter(
                teachers=teacher,
                classroom_type="scheduled",
                start_time__lt=current_datetime,
            ).aggregate(Sum("duration"))["duration__sum"]
            or 0,
            "active_students": teacher.students.filter(is_confirmed=True).count(),
            "inactive_students": teacher.students.filter(is_confirmed=False).count(),
            "avg_assignments_per_student": (
                round(total_assignments / total_students, 2) if total_students else 0
            ),
            "average_students_per_group": average_students_per_group["avg_students"],
            "total_assignments": total_assignments,
            "total_documents": Resource.objects.filter(
                assignment_links__assignment__teachers=teacher
            )
            .distinct()
            .count(),
            "total_class_groups": total_class_groups,
        }

        return Response({"data": stats}, status=status.HTTP_200_OK)

    except Teacher.DoesNotExist:
        return Response(
            {"error": "Teacher not found"}, status=status.HTTP_404_NOT_FOUND
        )


@api_view(["POST"])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def class_report(request):
    class_id = request.data.get("classID")
    if class_id is None:
        return Response(
            {"error": "classID is required"}, status=status.HTTP_400_BAD_REQUEST
        )
    try:
        class_event = ClassEvent.objects.get(id=class_id)
    except ClassEvent.DoesNotExist:
        return Response(
            {"error": "Class event not found"}, status=status.HTTP_404_NOT_FOUND
        )

    student_name = class_event.students.first()
    time_of_class = class_event.start_time
    subject = class_event.subject.name
    duration = class_event.duration
    class_summary = request.data.get("lesson_summary", "No summary provided")

    context = f"""
    You are a teacher's digital assistant. After the teacher conducts each class, you should receive a summary of it and its content
    and then use that summary to create a formal report. You should also set a suitable 10-minute homework task based on the lesson contents.
    Your answer should contain two sections with heading as follows:
    Summary: <summary>
    Homework: <homework task>

    Student Name: {student_name}
    Time of Class: {time_of_class}
    Subject: {subject}
    Duration: {duration}
    Teacher's summary: {class_summary}
    """

    client = OpenAI()
    completion = client.chat.completions.create(
        model="gpt-3.5-turbo", messages=[{"role": "system", "content": context}]
    )
    summary_start = completion.choices[0].message.content.find("Summary:")
    homework_start = completion.choices[0].message.content.find("Homework:")
    summary = (
        completion.choices[0]
        .message.content[summary_start + len("Summary:"): homework_start]
        .strip()
    )
    homework = (
        completion.choices[0]
        .message.content[homework_start + len("Homework:"):]
        .strip()
    )

    from docx import Document

    document = Document()
    document.add_heading(f"{student_name}'s Class Report", 0)
    table = document.add_table(rows=5, cols=1)
    table.rows[0].cells[0].text = f"Course: {subject}"
    table.rows[1].cells[0].text = f"Date: {time_of_class}"
    table.rows[2].cells[0].text = f"Duration: {duration} minutes"
    table.rows[3].cells[0].text = f"Class Summary: {summary}"
    table.rows[4].cells[0].text = f"Homework: {homework}"
    document.save("demo.docx")

    return JsonResponse({"message": completion.choices[0].message.content})


class SessionFeedbackViewSet(viewsets.ViewSet):
    authentication_classes = [TokenAuthentication]
    permission_classes = [IsAuthenticated]

    def create(self, request):
        user = request.user.get_real_instance()
        if not isinstance(user, Student):
            return Response(
                {"error": "Only students can submit session feedback"},
                status=status.HTTP_403_FORBIDDEN,
            )

        class_event_id = request.data.get("class_event")
        rating = request.data.get("rating")
        comment = request.data.get("comment", "")

        if class_event_id is None or rating is None:
            return Response(
                {"error": "class_event and rating are required"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            rating = int(rating)
        except (TypeError, ValueError):
            return Response(
                {"error": "rating must be an integer"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        if not 1 <= rating <= 5:
            return Response(
                {"error": "rating must be between 1 and 5"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            class_event = ClassEvent.objects.get(id=class_event_id)
        except ClassEvent.DoesNotExist:
            return Response(
                {"error": "Class event not found"}, status=status.HTTP_404_NOT_FOUND
            )

        if not class_event.students.filter(id=user.id).exists():
            return Response(
                {"error": "You were not enrolled in this class"},
                status=status.HTTP_403_FORBIDDEN,
            )

        try:
            with transaction.atomic():
                feedback = SessionFeedback.objects.create(
                    class_event=class_event,
                    student=user,
                    rating=rating,
                    comment=comment,
                )
            return Response(
                SessionFeedbackSerializer(feedback).data,
                status=status.HTTP_201_CREATED,
            )
        except IntegrityError:
            return Response(
                {"error": "You have already submitted feedback for this class"},
                status=status.HTTP_400_BAD_REQUEST,
            )

    @action(detail=True, methods=["get"])
    def aggregate(self, request, pk=None):
        user = request.user.get_real_instance()
        if not isinstance(user, Teacher):
            return Response(
                {"error": "Only teachers can view aggregated feedback"},
                status=status.HTTP_403_FORBIDDEN,
            )

        try:
            class_event = ClassEvent.objects.get(id=pk)
        except ClassEvent.DoesNotExist:
            return Response(
                {"error": "Class event not found"}, status=status.HTTP_404_NOT_FOUND
            )

        if not class_event.teachers.filter(id=user.id).exists():
            return Response(
                {"error": "You do not teach this class"},
                status=status.HTTP_403_FORBIDDEN,
            )

        feedbacks = SessionFeedback.objects.filter(class_event=class_event)
        agg = feedbacks.aggregate(
            average_rating=Avg("rating"),
            total_responses=Count("id"),
        )
        rating_distribution = {
            str(i): feedbacks.filter(rating=i).count() for i in range(1, 6)
        }
        comments = list(
            feedbacks.exclude(comment="").values(
                "student__first_name",
                "student__last_name",
                "comment",
                "rating",
                "created_at",
            )
        )
        return Response(
            {
                "class_event_id": pk,
                "average_rating": (
                    round(agg["average_rating"], 2) if agg["average_rating"] else None
                ),
                "total_responses": agg["total_responses"],
                "rating_distribution": rating_distribution,
                "comments": comments,
            }
        )


@api_view(["GET"])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def validate_classroom_access(request, access_token):
    try:
        classroom = ClassEvent.objects.get(access_token=access_token, is_active=True)

        if classroom.is_expired():
            return Response(
                {"error": "This classroom has expired", "expired": True},
                status=status.HTTP_410_GONE,
            )

        user = request.user.get_real_instance()
        if not classroom.can_access(user):
            return Response(
                {"error": "You do not have permission to access this classroom"},
                status=status.HTTP_403_FORBIDDEN,
            )

        serializer = ClassEventSerializer(classroom)
        user_id = user.id if hasattr(user, "id") else user.pk
        is_teacher = classroom.teachers.filter(id=user_id).exists()
        return Response(
            {
                "access_granted": True,
                "classroom": serializer.data,
                "user_role": "teacher" if is_teacher else "student",
            }
        )

    except ClassEvent.DoesNotExist:
        return Response(
            {"error": "Invalid classroom access token"},
            status=status.HTTP_404_NOT_FOUND,
        )


@api_view(["POST"])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def create_practice_classroom(request):
    user = request.user.get_real_instance()

    if not isinstance(user, Teacher):
        return Response(
            {"error": "Only teachers can create practice classrooms"},
            status=status.HTTP_403_FORBIDDEN,
        )

    try:
        from apps.subjects.models import Subject

        subject = (
            Subject.objects.filter(classevent__teachers=user)
            .annotate(usage_count=Count("classevent"))
            .order_by("-usage_count")
            .first()
        )
        if not subject:
            subject = Subject.objects.first()
            if not subject:
                subject = Subject.objects.create(name="Practice Subject")
    except Exception as e:
        return Response(
            {"error": f"Error creating practice classroom: {str(e)}"},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )

    practice_classroom = ClassEvent.objects.create(
        name=f"Practice Classroom - {user.first_name}",
        start_time=timezone.now(),
        duration=120,
        subject=subject,
        classroom_type="practice",
    )
    practice_classroom.teachers.add(user)

    serializer = ClassEventSerializer(practice_classroom)
    return Response(
        {
            "message": "Practice classroom created successfully",
            "classroom": serializer.data,
            "access_token": practice_classroom.access_token,
        },
        status=status.HTTP_201_CREATED,
    )


@api_view(["POST"])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def cleanup_expired_classrooms(request):
    user = request.user.get_real_instance()

    if not (user.is_staff or user.is_superuser):
        return Response(
            {"error": "Only administrators can trigger classroom cleanup"},
            status=status.HTTP_403_FORBIDDEN,
        )

    deleted_count = 0
    for classroom in ClassEvent.objects.filter(is_active=True):
        if classroom.is_expired():
            classroom.is_active = False
            classroom.save()
            deleted_count += 1

    return Response(
        {
            "message": f"Successfully deactivated {deleted_count} expired classrooms",
            "deactivated_count": deleted_count,
        }
    )
