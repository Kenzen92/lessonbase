import uuid
from django.shortcuts import render
from apps.storage.storage_backends import GridFSStorage
from faker import Faker
import logging
logger = logging.getLogger(__name__)

from django.shortcuts import render
from django.contrib.auth import login
from rest_framework.decorators import api_view, permission_classes, authentication_classes
from rest_framework.authentication import TokenAuthentication
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.response import Response
from rest_framework import permissions, status, generics
from rest_framework.views import APIView
from rest_framework.authtoken.models import Token
from .models import Chat, Message
from apps.classes.models import ClassEvent, Homework, TeachingResource
from apps.user_accounts.models import CustomerAccount, Staff, Student, Teacher
from apps.subjects.models import Subject
from .serializers import LoginSerializer, CustomerAccountSerializer, StudentSerializer, SubjectSerializer, TeacherSerializer, ChatSerializer, MessageSerializer
from apps.classes.serialisers import AssignHomeworkSerializer, ClassEventSerializer, HomeworkSerializer
from datetime import datetime
from django.core.mail import send_mail
from django.http import HttpResponse
from django.conf import settings
from django.template.loader import render_to_string
from django.http import JsonResponse
from openai import OpenAI
from docx import Document
from django.core.exceptions import ValidationError

@api_view(['GET'])
@permission_classes((permissions.AllowAny,))
def getData(request):
    items = ClassEvent.objects.all()
    serializer = ClassEventSerializer(items, many=True)
    return Response(serializer.data)

@api_view(['GET'])
@permission_classes((permissions.AllowAny,))
def createData(request):
    # Create 5 new teachers using faker
    fake = Faker()
    for _ in range(5):
        teacher = Teacher.objects.create(
            username=fake.user_name(),
            password=fake.password(),
            email=fake.email(),
            first_name=fake.first_name(),
            last_name=fake.last_name(),
            hire_date=fake.date_this_decade(),
        )

    # Create 5 new students using faker
    for _ in range(5):
        student = Student.objects.create(
            username=fake.user_name(),
            password=fake.password(),
            email=fake.email(),
            first_name=fake.first_name(),
            last_name=fake.last_name(),
            enrollment_date=fake.date_this_decade(),
        )

    # Create 20 class events using faker
    for _ in range(20):
        class_event = ClassEvent.objects.create(
            name=fake.text(max_nb_chars=50),
            start_time=fake.date_time(),
            duration=fake.random_int(min=1, max=6),
            subject=Subject.objects.order_by('?').first()
        )
        # Add random students and teachers to the class event
        for _ in range(fake.random_int(min=1, max=5)):  # Add 1 to 5 random students
            class_event.students.add(Student.objects.order_by('?').first())
        for _ in range(fake.random_int(min=1, max=3)):  # Add 1 to 3 random teachers
            class_event.teachers.add(Teacher.objects.order_by('?').first())

    return Response(200)


@api_view(['POST'])
@permission_classes((permissions.AllowAny,))
def userRegister(request):
    print("Request: ", request.data)
    # grab the ID's from the subjects and place them back into the request for the serializer to process
    subject_dicts = request.data['subjects']
    subject_ids = []
    for each_subject in subject_dicts:
        subject_ids.append(each_subject['value'])
    request.data['subjects'] = subject_ids
    serializer = CustomerAccountSerializer(data=request.data)
    if serializer.is_valid():
        user_type = serializer.validated_data['user_type']
        # Create an instance of the appropriate subclass
        if user_type == 1:
            user = Teacher.objects.create(username=serializer.validated_data['username'], hire_date=datetime.now())
            subject_ids = serializer.validated_data.get('subjects', [])  # Get list of subject IDs from request data
            user.subjects.add(*subject_ids)  # Add subjects to the user
        elif user_type == 2:
            user = Student.objects.create(username=serializer.validated_data['username'])
        else:
            user = Staff.objects.create(username=serializer.validated_data['username'])
        # Set other common fields (e.g., password)
        user.set_password(serializer.validated_data['password'])
        user.save()
        # Exclude subjects field from response
        serializer_data = serializer.validated_data.copy()
        serializer_data.pop('subjects', None)
        return Response(serializer_data, status=status.HTTP_201_CREATED)
    print(serializer.errors )
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)



@api_view(['POST'])
@permission_classes((permissions.AllowAny,))
def login(request):
    serializer = LoginSerializer(data=request.data)
    if serializer.is_valid(raise_exception=True):
        user = serializer.validated_data  # This should give you the user instance
        token, created = Token.objects.get_or_create(user=user)
        
        if user.polymorphic_ctype.model == "teacher":
            response_serializer = TeacherSerializer(user)
            return Response({"token": f"{token}", 'user': response_serializer.data}, status=status.HTTP_200_OK)
        else:
            response_serializer = StudentSerializer(user)
            return Response({"token": f"{token}", 'user': response_serializer.data}, status=status.HTTP_200_OK)
    else:
        logger.debug(serializer.errors)
        return Response({"error": f"{serializer.errors}"}, status=status.HTTP_400_BAD_REQUEST)
    

@api_view(['POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def logout(request):
    user = request.user
    try:
        # Attempt to retrieve the token associated with the user
        token = Token.objects.get(user=user)
        # Delete the token
        token.delete()
        return Response({"message": "Logout successful"}, status=status.HTTP_200_OK)
    except Token.DoesNotExist:
        # If the token doesn't exist, consider the user as already logged out
        return Response({"error": "User is not logged in"}, status=status.HTTP_400_BAD_REQUEST)


@api_view(['GET'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def students(request):
    queryset = Student.objects.all()
    serializer = StudentSerializer(queryset, many=True)
    return Response(serializer.data)

@api_view(['GET'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def teachers(request):
    queryset = Teacher.objects.all()
    serializer = TeacherSerializer(queryset, many=True)
    return Response(serializer.data)

@api_view(['GET'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def students_for_teacher(request):
    queryset = Student.objects.filter(teacher=request.user.id)
    serializer = StudentSerializer(queryset, many=True)
    return Response(serializer.data)

@api_view(['POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def connect_student_teacher(request):
    teacher = request.user.get_real_instance()
    student_ids = request.data.get('students', [])
    students = Student.objects.filter(id__in=student_ids)
    teacher.students.add(*students)
    teacher.save()
    return Response(status=status.HTTP_201_CREATED)
    
    

@api_view(['GET', 'POST', 'DELETE'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def class_events(request, class_id=None):
    if request.method == 'GET':
        user = request.user
        class_events = ClassEvent.objects.filter(students=user) | ClassEvent.objects.filter(teachers=user)
        serializer = ClassEventSerializer(class_events, many=True)
        return Response(serializer.data)

    
    elif request.method == 'DELETE':
        try:
            class_event = ClassEvent.objects.get(id=class_id)
            class_event.delete()
            return Response({'message': 'Class event deleted successfully'}, status=status.HTTP_204_NO_CONTENT)
        except ClassEvent.DoesNotExist:
            return Response({'error': 'Class event not found'}, status=status.HTTP_404_NOT_FOUND)
    
    elif request.method == 'POST':
        teacher_ids = [request.user.pk]
        student_ids = request.data.get('students', [])

        teachers = Teacher.objects.filter(pk__in=teacher_ids)
        students = Student.objects.filter(pk__in=student_ids)

        # subject = request.data.pop('subject')
        # print(subject)
        # try:
        #     subject_instance = Subject.objects.get(id=subject['id'])
        # except Subject.DoesNotExist:
        #     return Response({'message': 'Subject found'}, status=status.HTTP_404_NOT_FOUND)

        # request.data['subject'] = subject_instance
        serializer = ClassEventSerializer(data=request.data)
        if serializer.is_valid():
            class_event = serializer.save()
            class_event.teachers.set(teachers)
            class_event.students.set(students)
            return Response({"message": "Class event created successfully"}, status=status.HTTP_201_CREATED)
        else:
            print(serializer.errors)
            return Response({"error": serializer.errors}, status=status.HTTP_400_BAD_REQUEST)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    else:
        return Response({'message': 'Method not allowed'}, status=status.HTTP_405_METHOD_NOT_ALLOWED)


@api_view(['GET', 'POST'])    
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def subjects(request):
    user = request.user.get_real_instance()
    if request.method == 'POST':
        subjects_ids = request.data.get('subjects', [])
        try:    
            # If subjects array is empty, remove all subjects from the user
            if not subjects_ids:
                user.subjects.clear()
                return Response({'message': 'All subjects removed from the user'}, status=status.HTTP_200_OK)
            
            # Filter subjects to retain only those present in the request data
            valid_subjects = Subject.objects.filter(id__in=subjects_ids)
            
            # Remove subjects not present in the request data
            user.subjects.remove(*user.subjects.exclude(id__in=valid_subjects.values_list('id', flat=True)))
            
            # Add subjects not already connected to the user
            user.subjects.add(*valid_subjects.exclude(id__in=user.subjects.values_list('id', flat=True)))
            
            return Response({'message': 'User subjects updated successfully'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
    else:
        user = user.get_real_instance()
        subjects = user.subjects.all()
        serializer = SubjectSerializer(instance=subjects, many=True)  # Use 'instance' instead of 'data'
        return Response(serializer.data)
    

@api_view(['GET'])    
@permission_classes([AllowAny])
def all_subjects(request):
    subjects = Subject.objects.all()
    serializer = SubjectSerializer(instance=subjects, many=True)  # Use 'instance' instead of 'data'
    return Response(serializer.data)


@api_view(['GET', 'POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def profile(request):
    if request.method == 'GET':
        user = request.user.get_real_instance()
        if user.polymorphic_ctype.name == "teacher":
            serializer = TeacherSerializer(instance=user)
        else:
            serializer = StudentSerializer(instance=user)
        
        return Response(serializer.data)
    
    elif request.method == 'POST':
        user = request.user.get_real_instance()
        files = request.FILES
        
        # Adjust serializer to handle files
        if user.polymorphic_ctype.name == "teacher":
            serializer = TeacherSerializer(instance=user, data=request.data)
        else:
            serializer = StudentSerializer(instance=user, data=request.data)
        
        if serializer.is_valid():
            # Save the user instance
            user = serializer.save()
            
            # Handle file uploads and save to GridFS
            for file_key in files:
                uploaded_file = files[file_key]
                
                # Create a GridFSStorage instance
                gridfs_storage = GridFSStorage()  # Pass parameters if necessary
                
                # Save the file to GridFS
                gridfs_storage._save(uploaded_file.name, uploaded_file, context="profile_picture")
            
            return Response(serializer.data)
        else:
            print(serializer.errors)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        

@api_view(['POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def new_student(request):
    destination_email = request.data.get('email') # not using for now
    print(destination_email)
    sender_email = settings.DEFAULT_FROM_EMAIL
    # generate a token that is passed in the link url which will be used to validate the request to activate the account 
    token = uuid.uuid4()
    # Create an empty login account with default values, but attach the confirmation token.
    # student = Student.objects.filter(email=destination_email).first()
    student = None
    if student is None:  # If student doesn't exist, create a new one
        student = Student.objects.create(email=destination_email, username=uuid.uuid4()) 
    student.confirmation_token=token # replace any previous token
    teacher = request.user.get_real_instance()
    teacher.students.add(student)
    student.save()
    # No need to specify 'templates' in the path
    msg_plain = render_to_string('invitation_email.txt', {'token': token})
    msg_html = render_to_string('invitation_email.html', {'token': token})
    subject = 'Here is your invitation to join Kennysolutions'
    
    send_mail(
        subject,
        msg_plain,
        sender_email,
        ['jamespeterkenny@gmail.com'],
        html_message=msg_html,
        fail_silently=False,
    )
    return Response({"Message": 'Email sent successfully!'}, status=status.HTTP_200_OK)



@api_view(['GET', 'POST'])
@permission_classes((permissions.AllowAny,))
def confirm_account(request):
    if request.method == 'GET':
        token = request.GET.get('token')
        try:
            account = CustomerAccount.objects.get(confirmation_token=token)
        except CustomerAccount.DoesNotExist:
            return Response({"error": "Invalid token"}, status=status.HTTP_400_BAD_REQUEST)
       
        return render(request, "student_account_signup.html", {"token": token, confirm_account: "http://localhost:8000/confirm_account/"})
    elif request.method == 'POST':
        if request.data['username'] and request.data['password1']:
            token=request.data.get('token')
            try:
                account = CustomerAccount.objects.get(confirmation_token=token)
            except CustomerAccount.DoesNotExist:
                return Response({"error": "Invalid token"}, status=status.HTTP_400_BAD_REQUEST)
            
            account.username=request.data['username']
            account.set_password(request.data['password1'])
            account.save()
            account.confirmation_token = None # clear the token
            # Perform any additional actions here like redirecting to a success page, etc.
            return render(request, "account_confirmation.html")
        else:
            return Response({"error": "No username or password"}, status=status.HTTP_400_BAD_REQUEST)
        


class ChatListCreateView(generics.ListCreateAPIView):
    serializer_class = ChatSerializer
    authentication_classes = [TokenAuthentication]
    permission_classes = [permissions.AllowAny]

    def get_queryset(self):
        return Chat.objects.filter(participants=self.request.user)

    def perform_create(self, serializer):
        chat = serializer.save()
        user_ids = self.request.data.get('participants', [])  # Ensure this is a list
        print(f"User IDs to add: {user_ids}")

        # Add the current user to the chat
        chat.participants.add(self.request.user)
        
        # Add other participants
        if user_ids:
            for each_id in user_ids:
                try:
                    other_user = CustomerAccount.objects.get(id=each_id)
                    print("Added the other user yo: ", other_user)
                    chat.participants.add(other_user)
                except CustomerAccount.DoesNotExist:
                    print(f"User with ID {each_id} does not exist ")
        
        chat.save()

        # Debug prints
        print(f"Chat created with participants: {[participant.id for participant in chat.participants.all()]}")
        print(f"Chat participants (excluding current user): {[participant.id for participant in chat.participants.exclude(id=self.request.user.id)]}")

        # Return success response
        return Response(serializer.data, status=status.HTTP_201_CREATED)

class MessageListCreateView(generics.ListCreateAPIView):
    serializer_class = MessageSerializer
    authentication_classes = [TokenAuthentication]
    permission_classes = [permissions.AllowAny]
    
    def get_queryset(self):
        chat_id = self.kwargs['chat_id']
        return Message.objects.filter(chat_id=chat_id, chat__participants=self.request.user)

    def perform_create(self, serializer):
        chat_id = self.kwargs['chat_id']
        print("Creating chat with id: ", chat_id )
        serializer.save(sender=self.request.user, chat_id=chat_id)


@api_view(['POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def class_material(request):
    class_id = request.data.get('classID')
    if class_id is None:
        return Response({"error": "classID is required"}, status=status.HTTP_400_BAD_REQUEST)
    
    try:
        class_event = ClassEvent.objects.get(id=class_id)
    except ClassEvent.DoesNotExist:
        return Response({"error": "Class event not found"}, status=status.HTTP_404_NOT_FOUND)
    
    # Extract file and URL from the request
    uploaded_file = request.FILES.get('file', None)
    url = request.data.get('url', None)
    name = request.data.get('name', 'Unnamed Resource')
    description = request.data.get('description', '')
    subject_id = request.data.get('subjectID')

    # Check if subject exists
    try:
        subject = Subject.objects.get(id=subject_id)
    except Subject.DoesNotExist:
        return Response({"error": "Subject not found"}, status=status.HTTP_404_NOT_FOUND)

    # Validate that either a file or URL is provided, but not both
    if uploaded_file and url:
        return Response({"error": "Please provide only a file or a URL, not both."}, status=status.HTTP_400_BAD_REQUEST)
    if not uploaded_file and not url:
        return Response({"error": "Either a file or a URL must be provided."}, status=status.HTTP_400_BAD_REQUEST)

    # Handle file upload and save to GridFS if a file is provided
    file_url = None
    if uploaded_file:
        # Create a GridFSStorage instance
        gridfs_storage = GridFSStorage()
        # Save the file to GridFS
        file_name = gridfs_storage._save(uploaded_file.name, uploaded_file, context='teaching_resource')
        file_url = gridfs_storage.url(file_name)


    
    # Create a new TeachingResource instance
    teaching_resource = TeachingResource(
        name=name,
        description=description,
        file_url=file_url,
        url=url,
        subject=subject
    )
    
    try:
        teaching_resource.clean()  # Validate the instance
        teaching_resource.save()
    except ValidationError as e:
        return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    return Response({"message": "Teaching resource created successfully"}, status=status.HTTP_201_CREATED)

@api_view(['POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def class_report(request):
    class_id = request.data.get('classID')  # Assuming the class ID is passed as a query parameter
    if class_id is None:
        return Response({"error": "classID is required"}, status=status.HTTP_400_BAD_REQUEST)
    try:
        class_event = ClassEvent.objects.get(id=class_id)
    except ClassEvent.DoesNotExist:
        return Response({"error": "Class event not found"}, status=status.HTTP_404_NOT_FOUND)

    # Extract necessary information
    student_name = class_event.students.first()
    time_of_class = class_event.start_time
    subject = class_event.subject.name
    duration = class_event.duration

    # Prepare the context for the LLM
    class_summary = request.data.get("lesson_summary")

    # Prepare the context for the LLM
    class_summary = request.data.get("lesson_summary", "No summary provided")

   
    context = f"""
    You are a teacher's digital assistant. After the teacher conducts each class, you should receive a summary of it and its content
    and then use that summary to create a formal report. You should also set a suitable 10-minute homework task based on the lesson contents.
    Your answer should contain two sections with heading as follows:
    Summary: <summary>
    Homework: <homework task>

    Student Name: {student_name}
    Time of Class: {time_of_class}
    Subject: {subject}
    Duration: {duration}
    Teacher's summary: {class_summary}  
    """

    # Initialize OpenAI client and request a completion
    client = OpenAI()
    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": context}
        ]
    )
    # Split the text into summary and homework sections
    summary_start = completion.choices[0].message.content.find("Summary:")
    homework_start = completion.choices[0].message.content.find("Homework:")

    summary = completion.choices[0].message.content[summary_start + len("Summary:"):homework_start].strip()
    homework = completion.choices[0].message.content[homework_start + len("Homework:"):].strip()

    # Create a new document
    document = Document()
    document.add_heading(f"{student_name}'s Class Report", 0)

    table = document.add_table(rows=5, cols=1)
    table.rows[0].cells[0].text = f"Course: {subject}"
    table.rows[1].cells[0].text = f"Date: {time_of_class}"
    table.rows[2].cells[0].text = f"Duration: {duration} minutes"
    table.rows[3].cells[0].text = f"Class Summary: {summary}"
    table.rows[4].cells[0].text = f"Homework: {homework}"


    # Save the document
    document.save('demo.docx')

    # Extract the message content correctly
    message_content = completion.choices[0].message.content
    return JsonResponse({"message": message_content})


@api_view(['POST', 'GET'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def homework(request):
    user = request.user

    if request.method == "GET":
        # Get the homework ID from the query parameters
        homework_id = request.query_params.get('id', None)

        if homework_id:
            try:
                # Retrieve specific homework by ID
                homework_instance = Homework.objects.get(id=homework_id, teachers=user)
                serializer = AssignHomeworkSerializer(homework_instance)
                return Response(serializer.data, status=status.HTTP_200_OK)
            except Homework.DoesNotExist:
                return Response({"error": "Homework not found or you do not have access to it."}, status=status.HTTP_404_NOT_FOUND)

        else:
            # Retrieve all homework tasks assigned to the authenticated user
            homework_list = Homework.objects.filter(teachers=user)
            print(homework_list )
            serializer = HomeworkSerializer(homework_list, many=True)
            return Response(serializer.data, status=status.HTTP_200_OK)
        
    else:
        serializer = AssignHomeworkSerializer(data=request.data)
        if serializer.is_valid():
            # Save the instance
            homework_instance = serializer.save()

            # Optionally, add the user who is creating the homework to the teachers
            homework_instance.teachers.add(user)
            print(homework_instance )
            return Response({"success": "Homework instance created"}, status=status.HTTP_201_CREATED)
        else:
            return Response({"error": serializer.errors}, status=status.HTTP_400_BAD_REQUEST)
        

